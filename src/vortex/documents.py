"""Document intelligence (Phase 7) - file resolution and text extraction.

Reads PDF/DOCX/XLSX/TXT files by voice command ("read my resume", "summarize
notes.docx", "what does budget.xlsx say about Q3"). `extract_text` returns
the FULL text - truncation (for the simple summarize path) or chunking+
embedding (for the RAG-backed question-answering path, see rag.py) is a
decision for the caller, not this module.

`extract_pages` (added 2026-08-16) is the provenance-preserving sibling of
`extract_text`: same file support, but returns a list of {'page', 'section',
'text'} entries instead of one flattened string, so rag.py can chunk each
page/section separately and attach real page numbers (or a section/sheet
name, for formats with no page concept) to every chunk it stores - see
rag.py's module docstring for how that flows into retrieval and the assembled
answer prompt.

Also added 2026-08-16: OCR fallback for scanned/image-only PDF pages (see
_ocr_available/_ocr_page below). Both `extract_text` and `extract_pages` use
it automatically for PDF pages whose native PyMuPDF text is empty or
near-empty - the signal that a page is a rasterized image with no text layer,
not a genuinely blank page (imperfect heuristic, documented at OCR_MIN_TEXT_CHARS
below, not claimed to be exact).
"""
import logging
import os
from pathlib import Path

SEARCH_DIRS = [Path.home() / 'Desktop', Path.home() / 'Documents', Path.home() / 'Downloads']
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xls', '.txt', '.md'}


def _bool_env(name, default):
    val = os.getenv(name)
    if val is None:
        return default
    return val.strip().lower() not in ('0', 'false', 'no', 'off')


# Mirrors VortexConfig.ocr_enabled/ocr_language name-for-name and default-for-
# default (same env-var-name duplication pattern rag.py already uses for its
# own Postgres/Qdrant/embed-model constants - see rag.py's module docstring -
# since neither module is wired to take a VortexConfig instance from app.py
# yet; this stays a plain module-level constant like rag.py's, not an import
# of config.py, to match that existing precedent rather than mix patterns).
OCR_ENABLED = _bool_env('VORTEX_OCR_ENABLED', True)
OCR_LANGUAGE = os.getenv('VORTEX_OCR_LANGUAGE', 'eng')

# A page with fewer than this many non-whitespace characters of *native*
# extracted text is treated as an OCR candidate. Not a perfect signal - a
# genuinely near-blank page (e.g. just a page number or a section divider)
# will also trip this and get OCR'd for nothing worse than a wasted OCR call
# (OCR on a blank image returns empty/near-empty text too, so the original
# native text - however sparse - is kept whenever OCR doesn't produce more).
# 20 chars was picked as "clearly not a real paragraph of body text" without
# being so low that a short real heading gets missed as a scan candidate.
OCR_MIN_TEXT_CHARS = 20

_log = logging.getLogger('vortex.documents')

# Cached across calls within a process: OCR (when it runs at all) is the
# expensive part of PDF extraction, and both extract_text() and extract_pages()
# independently re-extract the same file on every voice question against it
# (app.py has no text cache of its own) - re-running Tesseract on every
# question for the same scanned document would be wasteful. Keyed on mtime so
# an edited file is picked up rather than serving stale OCR output.
_pdf_pages_cache: dict[tuple, list] = {}

# Cached result of the one-time "is Tesseract actually usable" probe (pip
# package importable AND the tesseract binary on PATH) - checked once, not
# once per page, and logged once, not spammed.
_ocr_status = None


def _ocr_available():
    """True only if OCR can actually run: OCR_ENABLED, pytesseract importable,
    AND the Tesseract OCR binary itself found on PATH. The binary is a
    separate, non-pip-installable dependency - installing pytesseract alone
    does NOT make OCR work, so this checks for the binary explicitly rather
    than assuming its presence from the Python package alone."""
    global _ocr_status
    if _ocr_status is not None:
        return _ocr_status
    if not OCR_ENABLED:
        _log.info('OCR fallback disabled via VORTEX_OCR_ENABLED=false.')
        _ocr_status = False
        return False
    try:
        import pytesseract  # noqa: F401
    except ImportError:
        _log.warning(
            'OCR fallback unavailable: the pytesseract package is not installed '
            "(pip install 'vortex-ai[documents]'). Scanned/image-only PDF pages "
            'will be skipped - only whatever native text PyMuPDF found will be used.')
        _ocr_status = False
        return False
    import shutil
    if not shutil.which('tesseract'):
        _log.warning(
            'OCR fallback unavailable: pytesseract is installed but the Tesseract '
            'OCR binary itself was not found on PATH (it is a separate, non-pip '
            'binary install, not bundled with the Python package). Scanned/'
            'image-only PDF pages will be skipped - only whatever native text '
            'PyMuPDF found will be used. To enable OCR on Windows: install '
            'Tesseract (e.g. https://github.com/UB-Mannheim/tesseract/wiki) and '
            'ensure tesseract.exe is on PATH, then restart VORTEX.')
        _ocr_status = False
        return False
    _log.info('OCR fallback available (Tesseract found on PATH).')
    _ocr_status = True
    return True


def _ocr_page(page):
    """Rasterize one PyMuPDF page and OCR it. Returns '' on any OCR failure
    (never raises) - a failed OCR attempt should fall back to whatever native
    text existed, not crash extraction of the rest of the document."""
    import pytesseract
    from PIL import Image
    try:
        # 200 DPI: high enough for Tesseract to read normal body text
        # reliably without the multi-second-per-page cost of a print-quality
        # (300+ DPI) render - not benchmarked precisely, but each doubling of
        # DPI quadruples pixel count and OCR time along with it.
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes('RGB', (pix.width, pix.height), pix.samples)
        return pytesseract.image_to_string(img, lang=OCR_LANGUAGE)
    except Exception as e:
        _log.warning(f'OCR failed for a page (page kept as native-text-only): {e}')
        return ''


def _extract_pdf_pages(path):
    """Returns [{'page': 1-indexed int, 'section': None, 'text': str}, ...],
    one entry per PDF page, OCR'd in place where native text is near-empty
    and OCR is available."""
    try:
        cache_key = (path, os.path.getmtime(path))
    except OSError:
        cache_key = None
    if cache_key is not None and cache_key in _pdf_pages_cache:
        return _pdf_pages_cache[cache_key]

    import fitz  # PyMuPDF
    pages = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc):
            text = page.get_text()
            if len(text.strip()) < OCR_MIN_TEXT_CHARS and _ocr_available():
                ocr_text = _ocr_page(page)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            pages.append({'page': i + 1, 'section': None, 'text': text})

    if cache_key is not None:
        _pdf_pages_cache[cache_key] = pages
    return pages


def _extract_docx_sections(path):
    """Groups paragraphs under the nearest preceding Heading-styled paragraph.
    .docx has no fixed page concept (Word reflows pagination at render time,
    and python-docx doesn't expose it), so "section" (heading text) is the
    cheap, honest substitute for provenance here - not invented structure,
    just paragraph.style.name, which python-docx already exposes for free."""
    import docx
    doc = docx.Document(path)
    sections = []
    heading, buf = None, []

    def flush():
        joined = '\n'.join(buf).strip()
        if joined:
            sections.append({'page': None, 'section': heading, 'text': joined})

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name.startswith('Heading') and p.text.strip():
            flush()
            heading, buf = p.text.strip(), []
        elif p.text:
            buf.append(p.text)
    flush()
    if not sections:
        sections = [{'page': None, 'section': None, 'text': '\n'.join(p.text for p in doc.paragraphs)}]
    return sections


def _extract_xlsx_sections(path):
    """One section per sheet - sheet name is free, cheap provenance (unlike a
    page number, which xlsx has no equivalent of)."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    sections = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(' | '.join(cells))
        if rows:
            sections.append({'page': None, 'section': sheet.title, 'text': '\n'.join(rows)})
    return sections


def resolve_document(name):
    """Find a file matching `name` (with or without extension) in common folders."""
    name = name.strip().strip('"\'').lower()
    candidates = []
    for d in SEARCH_DIRS:
        if not d.is_dir():
            continue
        for f in d.iterdir():
            if not f.is_file() or f.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            stem, full = f.stem.lower(), f.name.lower()
            if name == full or name == stem or name in stem:
                candidates.append(f)
    if not candidates:
        return None
    exact = [c for c in candidates if c.stem.lower() == name]
    pool = exact or candidates
    return str(max(pool, key=lambda p: p.stat().st_mtime))


def extract_text(path):
    """Flat full-text extraction (unchanged signature/behavior for callers
    that only need the whole document as one string, e.g. the plain-summarize
    path in main.py). For PDFs this now includes OCR'd text for any page
    whose native text was near-empty - see _extract_pdf_pages."""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        text = '\n'.join(p['text'] for p in _extract_pdf_pages(path))
    elif ext == '.docx':
        import docx
        text = '\n'.join(p.text for p in docx.Document(path).paragraphs)
    elif ext in ('.xlsx', '.xls'):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(' | '.join(cells))
        text = '\n'.join(rows)
    else:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    return text


def extract_pages(path):
    """Provenance-preserving extraction: returns a list of
    {'page': int|None, 'section': str|None, 'text': str} entries.

    - PDF: one entry per page, 1-indexed, OCR-backed (see _extract_pdf_pages).
    - DOCX: one entry per Heading-delimited section (section=heading text).
    - XLSX/XLS: one entry per sheet (section=sheet name).
    - TXT/MD: a single entry with no page/section (no substructure to find).

    Used by rag.py to attach page/section metadata to each stored chunk.
    extract_text(path) above remains the flat-string equivalent for callers
    that don't need provenance.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return _extract_pdf_pages(path)
    if ext == '.docx':
        return _extract_docx_sections(path)
    if ext in ('.xlsx', '.xls'):
        return _extract_xlsx_sections(path)
    return [{'page': None, 'section': None, 'text': extract_text(path)}]


def build_document_prompt(text, question):
    return (
        f'Document contents:\n{text}\n\n'
        f'Based only on the document above, {question}\n'
        'Answer concisely in a couple of spoken sentences.'
    )
